#!/usr/bin/env python3
"""
toc_validate.py - Table of Contents Validation for DOCX and PDF files.

Checks DOCX and PDF files for TOC quality issues including missing TOC fields,
empty placeholders, heading style mismatches, page break issues, and more.
Also validates TOC consistency across DOCX→PDF conversions.

Usage:
    python3 toc_validate.py check-docx output.docx
    python3 toc_validate.py check-pdf output.pdf
    python3 toc_validate.py check-conversion input.docx output.pdf

Output:
    JSON to stdout with structure:
    {
        "pass": true/false,
        "source": "filename",
        "check_type": "docx-toc"|"pdf-toc"|"conversion-toc",
        "errors": [...],
        "warnings": [...],
        "info": [...]
    }

Exit codes:
    0 = pass (no errors)
    1 = fail (errors found)
    2 = script error (bad args, file not found, etc.)

Dependencies:
    - Standard library (zipfile, xml.etree.ElementTree, etc.)
    - pdfplumber (for PDF checks)
    - pikepdf (optional, for link annotation checks)
"""

import sys
import os
import json
import re
import zipfile
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any

# ---------------------------------------------------------------------------
# XML namespace constants
# ---------------------------------------------------------------------------
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

# Standard heading style names (case-insensitive comparison)
STANDARD_HEADING_STYLES = {
    'heading1', 'heading2', 'heading3', 'heading4',
    'heading 1', 'heading 2', 'heading 3', 'heading 4',
    # Some localized variants
    '1', '2', '3', '4',
}

# TOC keywords to search for in PDF text
TOC_KEYWORDS = ['目录', '目 录', '目  录', 'table of contents', 'contents']

# Hint phrases that should not leak into final PDF
HINT_PHRASES = [
    '提示：本目录通过域代码生成',
    '右键更新域',
    'Update Field',
    'right-click',
    'Tip: This table of contents',
]

# Hint text indicators for DOCX styling check
HINT_INDICATORS = ['提示', 'Tip:', 'Update Field', '更新域']

# Gray color values (hex, case-insensitive)
GRAY_COLORS = {'808080', '999999', 'a0a0a0', 'a5a5a5', 'b0b0b0', 'c0c0c0',
               '888888', '777777', '666666', 'aaaaaa', 'bbbbbb', 'cccccc',
               '909090', '959595', '9a9a9a', 'a8a8a8', 'b8b8b8'}


# ---------------------------------------------------------------------------
# Result helpers
# ---------------------------------------------------------------------------
def make_item(code: str, message: str, severity: str) -> Dict[str, str]:
    """Create a single result item."""
    return {"code": code, "message": message, "severity": severity}


def make_result(source: str, check_type: str, errors: List, warnings: List,
                info: List) -> Dict[str, Any]:
    """Build the final result dict."""
    return {
        "pass": len(errors) == 0,
        "source": source,
        "check_type": check_type,
        "errors": errors,
        "warnings": warnings,
        "info": info,
    }


# ---------------------------------------------------------------------------
# DOCX XML parsing helpers
# ---------------------------------------------------------------------------
def parse_docx_xml(docx_path: str) -> Optional[ET.Element]:
    """Extract and parse document.xml from a .docx file.

    Returns the root Element or None if extraction fails.
    """
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            with z.open('word/document.xml') as f:
                return ET.parse(f).getroot()
    except (zipfile.BadZipFile, KeyError, ET.ParseError):
        return None


def get_all_paragraphs(root: ET.Element) -> List[ET.Element]:
    """Return all w:p elements in document order."""
    return root.findall('.//' + _w('p'))


def _w(tag: str) -> str:
    """Shorthand for word namespace tag."""
    return '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + tag


def get_paragraph_text(para: ET.Element) -> str:
    """Extract concatenated text from all w:t elements in a paragraph."""
    texts = []
    for t in para.findall('.//' + _w('t')):
        if t.text:
            texts.append(t.text)
    return ''.join(texts)


def get_paragraph_style(para: ET.Element) -> Optional[str]:
    """Get the pStyle val from a paragraph, or None."""
    pPr = para.find(_w('pPr'))
    if pPr is None:
        return None
    pStyle = pPr.find(_w('pStyle'))
    if pStyle is None:
        return None
    return pStyle.get(_w('val'))


def is_heading_style(style_val: Optional[str]) -> bool:
    """Check if a style value is a standard heading style."""
    if style_val is None:
        return False
    lower = style_val.lower().strip()
    # Check direct matches
    if lower in STANDARD_HEADING_STYLES:
        return True
    # Check "Heading" prefix pattern (e.g. "Heading1", "Heading 2")
    if lower.startswith('heading'):
        return True
    # Numeric style IDs sometimes used for headings
    if lower in ('1', '2', '3', '4'):
        return True
    return False


def is_any_heading_style(style_val: Optional[str]) -> bool:
    """Check if a style looks like any heading (standard or custom with 'heading')."""
    if style_val is None:
        return False
    lower = style_val.lower().strip()
    return lower.startswith('heading') or lower in ('1', '2', '3', '4')


def is_standard_heading_style(style_val: Optional[str]) -> bool:
    """Check if a style is specifically a standard Heading1-4."""
    if style_val is None:
        return False
    lower = style_val.lower().strip()
    return lower in {'heading1', 'heading2', 'heading3', 'heading4',
                     'heading 1', 'heading 2', 'heading 3', 'heading 4'}


def paragraph_is_bold_large(para: ET.Element) -> bool:
    """Check if a paragraph has bold text and large font (≥28 half-points / 14pt).

    Checks both paragraph-level and run-level properties.
    """
    is_bold = False
    is_large = False

    # Check paragraph-level properties
    pPr = para.find(_w('pPr'))
    if pPr is not None:
        rPr = pPr.find(_w('rPr'))
        if rPr is not None:
            b = rPr.find(_w('b'))
            if b is not None:
                b_val = b.get(_w('val'))
                if b_val is None or b_val.lower() not in ('false', '0', 'off'):
                    is_bold = True
            sz = rPr.find(_w('sz'))
            if sz is not None:
                try:
                    size = int(sz.get(_w('val'), '0'))
                    if size >= 28:
                        is_large = True
                except (ValueError, TypeError):
                    pass

    # Check run-level properties
    for run in para.findall(_w('r')):
        rPr = run.find(_w('rPr'))
        if rPr is None:
            continue
        b = rPr.find(_w('b'))
        if b is not None:
            b_val = b.get(_w('val'))
            if b_val is None or b_val.lower() not in ('false', '0', 'off'):
                is_bold = True
        sz = rPr.find(_w('sz'))
        if sz is not None:
            try:
                size = int(sz.get(_w('val'), '0'))
                if size >= 28:
                    is_large = True
            except (ValueError, TypeError):
                pass

    return is_bold and is_large


def docx_has_toc_field(root: ET.Element) -> bool:
    """Check if the document has a TOC field code.

    Looks for:
    - <w:fldSimple> with w:instr containing "TOC"
    - <w:instrText> containing "TOC"
    """
    # Check fldSimple
    for fld in root.findall('.//' + _w('fldSimple')):
        instr = fld.get(_w('instr'), '')
        if 'TOC' in instr.upper():
            return True

    # Check instrText
    for instr in root.findall('.//' + _w('instrText')):
        if instr.text and 'TOC' in instr.text.upper():
            return True

    return False


def find_toc_field_boundaries(root: ET.Element) -> Tuple[Optional[ET.Element], Optional[ET.Element], Optional[ET.Element]]:
    """Find the TOC field begin/separate/end fldChar elements.

    Returns (begin_elem, separate_elem, end_elem) — any may be None.
    We search for the TOC instrText and then find the corresponding
    fldChar markers.
    """
    body = root.find(_w('body'))
    if body is None:
        return None, None, None

    all_paragraphs = list(body)  # Direct children of body

    # Flatten all runs across all paragraphs to find field structure
    # We need to track field nesting to find the right begin/separate/end
    in_toc_field = False
    toc_begin_para_idx = None
    toc_separate_para_idx = None
    toc_end_para_idx = None
    field_depth = 0

    for para_idx, elem in enumerate(all_paragraphs):
        if elem.tag != _w('p'):
            continue
        for run in elem.findall(_w('r')):
            # Check for instrText with TOC
            instr = run.find(_w('instrText'))
            if instr is not None and instr.text and 'TOC' in instr.text.upper():
                in_toc_field = True

            fldChar = run.find(_w('fldChar'))
            if fldChar is not None:
                fld_type = fldChar.get(_w('fldCharType'), '')
                if fld_type == 'begin':
                    field_depth += 1
                    if not in_toc_field and toc_begin_para_idx is None:
                        # Mark tentatively; will confirm when we see instrText
                        pass
                    if in_toc_field and toc_begin_para_idx is None:
                        toc_begin_para_idx = para_idx
                elif fld_type == 'separate':
                    if in_toc_field and toc_separate_para_idx is None:
                        toc_separate_para_idx = para_idx
                elif fld_type == 'end':
                    if in_toc_field and field_depth <= 1:
                        toc_end_para_idx = para_idx
                        in_toc_field = False
                    field_depth = max(0, field_depth - 1)

    return toc_begin_para_idx, toc_separate_para_idx, toc_end_para_idx


def find_toc_field_boundaries_v2(root: ET.Element) -> Dict[str, Any]:
    """Enhanced TOC boundary finder that works with nested fields.

    Returns dict with:
        'has_toc': bool
        'begin_para_idx': int or None
        'separate_para_idx': int or None
        'end_para_idx': int or None
        'toc_entry_texts': list of str (text between separate and end)
    """
    body = root.find(_w('body'))
    if body is None:
        return {'has_toc': False, 'begin_para_idx': None,
                'separate_para_idx': None, 'end_para_idx': None,
                'toc_entry_texts': []}

    paragraphs = [e for e in body if e.tag == _w('p')]

    # Phase 1: Find the TOC instrText and its surrounding begin marker
    toc_begin_idx = None
    toc_separate_idx = None
    toc_end_idx = None

    # Track all fldChar positions
    events = []  # (para_idx, event_type, element)
    for pi, para in enumerate(paragraphs):
        for run in para.findall('.//' + _w('r')):
            fldChar = run.find(_w('fldChar'))
            if fldChar is not None:
                events.append((pi, fldChar.get(_w('fldCharType'), ''), run))
            instr = run.find(_w('instrText'))
            if instr is not None and instr.text and 'TOC' in instr.text.upper():
                events.append((pi, 'toc_instr', run))

    # Find TOC field boundaries using field nesting
    depth = 0
    found_toc = False
    toc_depth = None

    begin_stack = []  # Stack of (para_idx, depth)

    for pi, evt, run in events:
        if evt == 'begin':
            depth += 1
            begin_stack.append((pi, depth))
        elif evt == 'toc_instr':
            if not found_toc and begin_stack:
                found_toc = True
                toc_begin_idx = begin_stack[-1][0]
                toc_depth = begin_stack[-1][1]
        elif evt == 'separate':
            if found_toc and toc_separate_idx is None and depth == toc_depth:
                toc_separate_idx = pi
        elif evt == 'end':
            if found_toc and toc_end_idx is None and depth == toc_depth:
                toc_end_idx = pi
            depth = max(0, depth - 1)
            if begin_stack:
                begin_stack.pop()

    # Phase 2: Extract TOC entry texts between separate and end
    toc_entry_texts = []
    if toc_separate_idx is not None and toc_end_idx is not None:
        for pi in range(toc_separate_idx, toc_end_idx + 1):
            if pi < len(paragraphs):
                text = get_paragraph_text(paragraphs[pi]).strip()
                if text:
                    toc_entry_texts.append(text)

    return {
        'has_toc': found_toc,
        'begin_para_idx': toc_begin_idx,
        'separate_para_idx': toc_separate_idx,
        'end_para_idx': toc_end_idx,
        'toc_entry_texts': toc_entry_texts,
    }


def check_toc_has_content(root: ET.Element, separate_para_idx: Optional[int],
                          end_para_idx: Optional[int]) -> bool:
    """Check if there are w:t elements between the separate and end markers.

    Looks at all paragraphs between the separate and end field char markers.
    """
    if separate_para_idx is None or end_para_idx is None:
        return False

    body = root.find(_w('body'))
    if body is None:
        return False

    paragraphs = [e for e in body if e.tag == _w('p')]

    for pi in range(separate_para_idx, min(end_para_idx + 1, len(paragraphs))):
        para = paragraphs[pi]
        for t in para.findall('.//' + _w('t')):
            if t.text and t.text.strip():
                return True
    return False


def fuzzy_match(text_a: str, text_b: str) -> bool:
    """Check if two strings match fuzzily.

    Match if one contains the other, or they share >60% of characters.
    """
    a = text_a.strip().lower()
    b = text_b.strip().lower()

    if not a or not b:
        return False

    # One contains the other
    if a in b or b in a:
        return True

    # Character overlap >60%
    set_a = set(a)
    set_b = set(b)
    if not set_a or not set_b:
        return False
    intersection = set_a & set_b
    union = set_a | set_b
    similarity = len(intersection) / len(union)
    return similarity > 0.6


def _detect_language(texts: list) -> str:
    """Detect the primary language of a list of text strings.

    Returns 'zh' if more than half contain Chinese characters, else 'en'.
    """
    if not texts:
        return 'en'
    total = len(texts)
    chinese_count = sum(1 for t in texts if re.search(r'[\u4e00-\u9fff]', t))
    return 'zh' if chinese_count > total / 2 else 'en'


def _get_heading_level(style_val: Optional[str]) -> int:
    """Extract heading level (1-9) from a style value. Returns 0 if not a heading."""
    if style_val is None:
        return 0
    lower = style_val.lower().strip()
    # "heading1", "heading 1", "heading2", etc.
    m = re.match(r'heading\s*(\d+)', lower)
    if m:
        return int(m.group(1))
    # Numeric style IDs
    if lower in ('1', '2', '3', '4', '5', '6', '7', '8', '9'):
        return int(lower)
    return 0


def check_run_hint_style(run: ET.Element) -> Tuple[bool, bool]:
    """Check if a run has gray color and small font size.

    Returns (has_gray_color, has_small_font).
    """
    has_gray = False
    has_small = False

    rPr = run.find(_w('rPr'))
    if rPr is None:
        return False, False

    color = rPr.find(_w('color'))
    if color is not None:
        val = color.get(_w('val'), '').lower()
        if val in GRAY_COLORS:
            has_gray = True
        # Also check if it's any gray-ish color (same R, G, B values or close)
        if len(val) == 6:
            try:
                r = int(val[0:2], 16)
                g = int(val[2:4], 16)
                b = int(val[4:6], 16)
                # Gray if R, G, B are all close to each other and in mid-range
                if (abs(r - g) < 30 and abs(g - b) < 30 and abs(r - b) < 30
                        and 80 <= r <= 210):
                    has_gray = True
            except ValueError:
                pass

    sz = rPr.find(_w('sz'))
    if sz is not None:
        try:
            size = int(sz.get(_w('val'), '0'))
            if size <= 18:  # 18 half-points = 9pt
                has_small = True
        except (ValueError, TypeError):
            pass

    return has_gray, has_small


# ---------------------------------------------------------------------------
# check-docx implementation
# ---------------------------------------------------------------------------
def check_docx(docx_path: str) -> Dict[str, Any]:
    """Run all DOCX TOC validation checks.

    Returns the result dict.
    """
    errors: List[Dict] = []
    warnings: List[Dict] = []
    info: List[Dict] = []
    source = os.path.basename(docx_path)

    # Parse document.xml
    root = parse_docx_xml(docx_path)
    if root is None:
        return make_result(source, "docx-toc",
                           [make_item("PARSE_ERROR",
                                      "Failed to parse DOCX file. File may be corrupted.",
                                      "error")],
                           [], [])

    body = root.find(_w('body'))
    if body is None:
        return make_result(source, "docx-toc", [], [],
                           [make_item("EMPTY_BODY",
                                      "Document body is empty.",
                                      "info")])

    paragraphs = [e for e in body if e.tag == _w('p')]

    # Detect TOC field boundaries
    toc_info = find_toc_field_boundaries_v2(root)
    has_toc = toc_info['has_toc']
    toc_begin_idx = toc_info['begin_para_idx']
    toc_separate_idx = toc_info['separate_para_idx']
    toc_end_idx = toc_info['end_para_idx']
    toc_entry_texts = toc_info['toc_entry_texts']

    # Also check for fldSimple-based TOC
    if not has_toc:
        for fld in root.findall('.//' + _w('fldSimple')):
            instr = fld.get(_w('instr'), '')
            if 'TOC' in instr.upper():
                has_toc = True
                break

    # Also check for SDT-wrapped TOC (e.g. generated by fix-docx)
    if not has_toc:
        if docx_has_toc_field(root):
            has_toc = True

    # Count headings (all paragraphs with heading styles)
    all_heading_paras = []
    for pi, para in enumerate(paragraphs):
        style = get_paragraph_style(para)
        if is_any_heading_style(style):
            all_heading_paras.append((pi, para, style))

    # Content headings: headings AFTER the TOC end (or all if no TOC)
    content_heading_paras = []
    if toc_end_idx is not None:
        for pi, para, style in all_heading_paras:
            if pi > toc_end_idx:
                content_heading_paras.append((pi, para, style))
    else:
        content_heading_paras = list(all_heading_paras)

    # ---- CHECK 1: TOC_FIELD_MISSING ----
    heading_count = len(all_heading_paras)
    if heading_count >= 3 and not has_toc:
        warnings.append(make_item(
            "TOC_FIELD_MISSING",
            f"Document has {heading_count} headings but no Table of Contents.",
            "warning"
        ))

    # ---- CHECK 2: TOC_PLACEHOLDER_EMPTY ----
    if has_toc and toc_separate_idx is not None and toc_end_idx is not None:
        has_content = check_toc_has_content(root, toc_separate_idx, toc_end_idx)
        if not has_content:
            errors.append(make_item(
                "TOC_PLACEHOLDER_EMPTY",
                "TOC field exists but has no placeholder entries. Run add_toc_placeholders.py.",
                "error"
            ))

    # ---- CHECK 3: TOC_HEADING_STYLE ----
    # Scan ALL paragraphs after TOC (not just those in content_heading_paras)
    # to catch bold+large paragraphs with non-heading styles that TOC won't see.
    if has_toc:
        start_idx = (toc_end_idx + 1) if toc_end_idx is not None else 0
        for pi in range(start_idx, len(paragraphs)):
            para = paragraphs[pi]
            style = get_paragraph_style(para)
            if paragraph_is_bold_large(para) and not is_standard_heading_style(style):
                text = get_paragraph_text(para).strip()
                if text:
                    truncated = text[:50] + ('...' if len(text) > 50 else '')
                    style_name = style if style else '(none)'
                    errors.append(make_item(
                        "TOC_HEADING_STYLE",
                        f"Paragraph '{truncated}' uses custom style '{style_name}' "
                        f"instead of HeadingLevel. TOC will not pick it up.",
                        "error"
                    ))

    # ---- CHECK 4: TOC_ENTRY_MISMATCH ----
    if toc_entry_texts and content_heading_paras:
        heading_texts = [get_paragraph_text(para).strip()
                         for _, para, _ in content_heading_paras
                         if get_paragraph_text(para).strip()]
        if heading_texts:
            unmatched = 0
            for ht in heading_texts:
                matched = any(fuzzy_match(ht, et) for et in toc_entry_texts)
                if not matched:
                    unmatched += 1

            match_ratio = (len(heading_texts) - unmatched) / len(heading_texts)
            if match_ratio < 0.5:
                errors.append(make_item(
                    "TOC_ENTRY_MISMATCH",
                    f"TOC placeholder entries don't match actual headings. "
                    f"{unmatched} of {len(heading_texts)} headings not found in TOC.",
                    "error"
                ))

    # ---- CHECK 5: TOC_NO_PAGEBREAK ----
    if toc_end_idx is not None:
        found_pagebreak = False
        # Check up to 2 paragraphs after TOC end
        check_end = min(toc_end_idx + 3, len(paragraphs))
        for pi in range(toc_end_idx, check_end):
            para = paragraphs[pi]
            # Check for <w:br w:type="page"/>
            for br in para.findall('.//' + _w('br')):
                if br.get(_w('type')) == 'page':
                    found_pagebreak = True
                    break
            # Check for <w:lastRenderedPageBreak/>
            if para.findall('.//' + _w('lastRenderedPageBreak')):
                found_pagebreak = True
            if found_pagebreak:
                break

        if not found_pagebreak:
            warnings.append(make_item(
                "TOC_NO_PAGEBREAK",
                "No page break found after TOC. Content may run into the table of contents.",
                "warning"
            ))

    # ---- CHECK 6: TOC_HINT_STYLE ----
    for para in paragraphs:
        text = get_paragraph_text(para).strip()
        has_hint = any(indicator in text for indicator in HINT_INDICATORS)
        if has_hint:
            # Check if runs containing hint text are properly styled
            properly_styled = True
            for run in para.findall(_w('r')):
                run_text = ''
                for t in run.findall(_w('t')):
                    if t.text:
                        run_text += t.text
                if any(ind in run_text for ind in HINT_INDICATORS):
                    has_gray, has_small = check_run_hint_style(run)
                    if not (has_gray and has_small):
                        properly_styled = False
                        break

            if not properly_styled:
                warnings.append(make_item(
                    "TOC_HINT_STYLE",
                    "TOC hint text found but not styled as gray/small. "
                    "It may look like regular content.",
                    "warning"
                ))
            break  # Only report once

    return make_result(source, "docx-toc", errors, warnings, info)


# ---------------------------------------------------------------------------
# check-pdf implementation
# ---------------------------------------------------------------------------
def check_pdf(pdf_path: str) -> Dict[str, Any]:
    """Run all PDF TOC validation checks.

    Returns the result dict.
    """
    errors: List[Dict] = []
    warnings: List[Dict] = []
    info: List[Dict] = []
    source = os.path.basename(pdf_path)

    try:
        import pdfplumber
    except ImportError:
        return make_result(source, "pdf-toc",
                           [make_item("DEPENDENCY_MISSING",
                                      "pdfplumber is not installed. Run: pip install pdfplumber",
                                      "error")],
                           [], [])

    try:
        pdf = pdfplumber.open(pdf_path)
    except Exception as e:
        return make_result(source, "pdf-toc",
                           [make_item("PARSE_ERROR",
                                      f"Failed to open PDF: {str(e)[:100]}",
                                      "error")],
                           [], [])

    total_pages = len(pdf.pages)
    if total_pages == 0:
        pdf.close()
        return make_result(source, "pdf-toc", [], [],
                           [make_item("EMPTY_PDF", "PDF has no pages.", "info")])

    # Extract text from first 5 pages (or all if <5)
    check_pages = min(5, total_pages)
    page_texts = {}
    for i in range(check_pages):
        try:
            text = pdf.pages[i].extract_text() or ''
        except Exception:
            text = ''
        page_texts[i] = text

    # ---- CHECK 1: TOC_NOT_FOUND ----
    toc_pages = []
    for page_idx, text in page_texts.items():
        text_lower = text.lower()
        for kw in TOC_KEYWORDS:
            if kw.lower() in text_lower:
                toc_pages.append(page_idx)
                break

    if not toc_pages and total_pages > 5:
        warnings.append(make_item(
            "TOC_NOT_FOUND",
            f"No TOC detected in first 5 pages of a {total_pages}-page document.",
            "warning"
        ))

    # ---- CHECK 1b: TOC_ON_FIRST_PAGE ----
    # If TOC appears on page 1, it likely means either:
    # (a) there is no cover page before the TOC, or
    # (b) the TOC and body content are not separated by a page break
    if toc_pages and 0 in toc_pages and total_pages > 1:
        errors.append(make_item(
            "TOC_ON_FIRST_PAGE",
            "TOC detected on page 1. A cover page should precede the TOC "
            "(expected structure: Cover → TOC → Content). "
            "Either the cover page is missing or the TOC was not separated by a page break.",
            "error"
        ))

    # ---- CHECK 2 & 3 & 4: TOC entry analysis ----
    # Regex to find lines where the last token is a number (page reference)
    entry_pattern = re.compile(r'^(.+?)\s+(\d{1,4})\s*$')

    toc_entries = []  # List of (title_text, page_number)
    if toc_pages:
        for page_idx in toc_pages:
            text = page_texts.get(page_idx, '')
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                m = entry_pattern.match(line)
                if m:
                    title = m.group(1).strip()
                    page_num = int(m.group(2))
                    if 1 <= page_num <= 9999 and title:
                        toc_entries.append((title, page_num))

    if toc_pages:
        # CHECK 2: TOC_NO_ENTRIES
        if len(toc_entries) < 2:
            errors.append(make_item(
                "TOC_NO_ENTRIES",
                "TOC page found but contains fewer than 2 entries.",
                "error"
            ))

        if toc_entries:
            # CHECK 3: TOC_PAGES_INVALID
            invalid_entries = []
            for title, page_num in toc_entries:
                if page_num < 1 or page_num > total_pages:
                    invalid_entries.append((title, page_num))

            for title, page_num in invalid_entries:
                truncated = title[:50] + ('...' if len(title) > 50 else '')
                errors.append(make_item(
                    "TOC_PAGES_INVALID",
                    f"TOC entry '{truncated}' references page {page_num} "
                    f"but document only has {total_pages} pages.",
                    "error"
                ))

            # CHECK 4: TOC_ALL_SAME_PAGE
            if len(toc_entries) >= 2:
                page_nums = set(pn for _, pn in toc_entries)
                if len(page_nums) == 1:
                    same_page = page_nums.pop()
                    errors.append(make_item(
                        "TOC_ALL_SAME_PAGE",
                        f"All TOC entries point to page {same_page}. "
                        f"This likely means placeholder page numbers were not updated.",
                        "error"
                    ))

    # ---- CHECK 5: TOC_LINKS_MISSING ----
    if toc_entries and toc_pages:
        has_links = False
        for page_idx in toc_pages:
            try:
                page = pdf.pages[page_idx]
                # Try annots (annotations)
                annots = page.annots
                if annots:
                    has_links = True
                    break
                # Try hyperlinks
                hyperlinks = page.hyperlinks
                if hyperlinks:
                    has_links = True
                    break
            except (AttributeError, Exception):
                pass

        if not has_links:
            # Also try pikepdf for more thorough annotation check
            try:
                import pikepdf
                pike_pdf = pikepdf.open(pdf_path)
                for page_idx in toc_pages:
                    if page_idx < len(pike_pdf.pages):
                        pike_page = pike_pdf.pages[page_idx]
                        if '/Annots' in pike_page:
                            annots = pike_page['/Annots']
                            if len(annots) > 0:
                                has_links = True
                                break
                pike_pdf.close()
            except (ImportError, Exception):
                pass

        if not has_links:
            warnings.append(make_item(
                "TOC_LINKS_MISSING",
                "TOC entries found but no clickable links detected.",
                "warning"
            ))

    pdf.close()
    return make_result(source, "pdf-toc", errors, warnings, info)


# ---------------------------------------------------------------------------
# check-conversion implementation
# ---------------------------------------------------------------------------
def check_conversion(docx_path: str, pdf_path: str) -> Dict[str, Any]:
    """Run DOCX→PDF conversion TOC consistency checks.

    Returns the result dict.
    """
    errors: List[Dict] = []
    warnings: List[Dict] = []
    info: List[Dict] = []
    source = f"{os.path.basename(docx_path)} → {os.path.basename(pdf_path)}"

    # Parse DOCX
    docx_root = parse_docx_xml(docx_path)
    if docx_root is None:
        return make_result(source, "conversion-toc",
                           [make_item("PARSE_ERROR",
                                      "Failed to parse source DOCX file.",
                                      "error")],
                           [], [])

    # Check DOCX has TOC
    docx_has_toc = docx_has_toc_field(docx_root)

    # Parse PDF
    try:
        import pdfplumber
    except ImportError:
        return make_result(source, "conversion-toc",
                           [make_item("DEPENDENCY_MISSING",
                                      "pdfplumber is not installed.",
                                      "error")],
                           [], [])

    try:
        pdf = pdfplumber.open(pdf_path)
    except Exception as e:
        return make_result(source, "conversion-toc",
                           [make_item("PARSE_ERROR",
                                      f"Failed to open PDF: {str(e)[:100]}",
                                      "error")],
                           [], [])

    total_pages = len(pdf.pages)

    # Extract all PDF text
    all_pdf_text = ''
    page_texts = {}
    for i in range(total_pages):
        try:
            text = pdf.pages[i].extract_text() or ''
        except Exception:
            text = ''
        page_texts[i] = text
        all_pdf_text += text + '\n'

    # Find TOC pages in PDF
    toc_pages = []
    check_pages = min(5, total_pages)
    for i in range(check_pages):
        text_lower = page_texts.get(i, '').lower()
        for kw in TOC_KEYWORDS:
            if kw.lower() in text_lower:
                toc_pages.append(i)
                break

    pdf_has_toc = len(toc_pages) > 0

    # ---- CHECK 1: CONV_TOC_LOST ----
    if docx_has_toc and not pdf_has_toc and total_pages > 5:
        errors.append(make_item(
            "CONV_TOC_LOST",
            "Source DOCX has TOC but converted PDF does not. "
            "TOC was lost during conversion.",
            "error"
        ))

    # ---- CHECK 2: CONV_HINT_LEAKED ----
    all_text_lower = all_pdf_text.lower()
    for phrase in HINT_PHRASES:
        if phrase.lower() in all_text_lower:
            # Find the actual matched text (up to 60 chars)
            idx = all_text_lower.index(phrase.lower())
            matched = all_pdf_text[idx:idx + len(phrase)]
            truncated = matched[:60] + ('...' if len(matched) > 60 else '')
            errors.append(make_item(
                "CONV_HINT_LEAKED",
                f"TOC hint text leaked into PDF: '{truncated}'. "
                f"Clean hints before conversion.",
                "error"
            ))
            break  # Report only the first match

    # ---- CHECK 3: CONV_HEADING_DRIFT ----
    # Count DOCX headings
    body = docx_root.find(_w('body'))
    docx_heading_count = 0
    if body is not None:
        for para in body:
            if para.tag != _w('p'):
                continue
            style = get_paragraph_style(para)
            if is_any_heading_style(style):
                docx_heading_count += 1

    # Count PDF TOC entries
    entry_pattern = re.compile(r'^(.+?)\s+(\d{1,4})\s*$')
    pdf_toc_entry_count = 0
    if toc_pages:
        for page_idx in toc_pages:
            text = page_texts.get(page_idx, '')
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                m = entry_pattern.match(line)
                if m:
                    page_num = int(m.group(2))
                    if 1 <= page_num <= 9999:
                        pdf_toc_entry_count += 1

    if docx_heading_count > 0 and pdf_toc_entry_count > 0:
        drift = abs(docx_heading_count - pdf_toc_entry_count)
        drift_pct = (drift / docx_heading_count) * 100
        if drift_pct > 30:
            warnings.append(make_item(
                "CONV_HEADING_DRIFT",
                f"DOCX has {docx_heading_count} headings but PDF TOC has "
                f"{pdf_toc_entry_count} entries ({drift_pct:.0f}% drift).",
                "warning"
            ))

    pdf.close()
    return make_result(source, "conversion-toc", errors, warnings, info)


# ---------------------------------------------------------------------------
# fix-docx implementation
# ---------------------------------------------------------------------------
def _find_toc_sdt_indices(body_elem) -> List[int]:
    """Find indices of SDT elements in body that contain TOC.

    Returns list of indices into body's direct children.
    """
    indices = []
    for idx, child in enumerate(body_elem):
        if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt':
            # Check if this SDT contains TOC-related content
            for instr in child.findall('.//' + _w('instrText')):
                if instr.text and 'TOC' in instr.text.upper():
                    indices.append(idx)
                    break
            else:
                # Also check alias/tag
                sdtPr = child.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr')
                if sdtPr is not None:
                    alias = sdtPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias')
                    if alias is not None and alias.get(_w('val'), '').upper() in ('TOC', '目录'):
                        indices.append(idx)
                        continue
                    docPartObj = sdtPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docPartObj')
                    if docPartObj is not None:
                        docPartGallery = docPartObj.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docPartGallery')
                        if docPartGallery is not None and 'toc' in docPartGallery.get(_w('val'), '').lower():
                            indices.append(idx)
    return indices


def _find_toc_field_para_range(body_elem) -> Tuple[Optional[int], Optional[int]]:
    """Find the range of paragraph indices that make up a TOC field code block.

    Returns (start_idx, end_idx) inclusive, or (None, None) if not found.
    These are indices into body's direct children.
    """
    children = list(body_elem)
    in_toc = False
    toc_depth = None
    depth = 0
    start_idx = None
    end_idx = None

    for ci, child in enumerate(children):
        if child.tag != _w('p'):
            continue
        for run in child.findall('.//' + _w('r')):
            instr = run.find(_w('instrText'))
            if instr is not None and instr.text and 'TOC' in instr.text.upper():
                in_toc = True

            fldChar = run.find(_w('fldChar'))
            if fldChar is not None:
                fld_type = fldChar.get(_w('fldCharType'), '')
                if fld_type == 'begin':
                    depth += 1
                    if in_toc and start_idx is None:
                        # The begin was before instrText; look back
                        start_idx = ci
                        toc_depth = depth
                    elif not in_toc and start_idx is None:
                        # tentative; may become TOC if instrText follows
                        pass
                elif fld_type == 'end':
                    if in_toc and depth == toc_depth:
                        end_idx = ci
                        in_toc = False
                    depth = max(0, depth - 1)

    # If we found instrText but start_idx wasn't set (begin was in the same para before instrText)
    # Re-scan more carefully
    if in_toc and start_idx is None:
        # Fall back to find_toc_field_boundaries_v2 style
        pass

    return start_idx, end_idx


def fix_docx(docx_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
    """Detect TOC issues in a DOCX and fix them, outputting a new DOCX file.

    Returns the result dict.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    errors: List[Dict] = []
    warnings: List[Dict] = []
    info_list: List[Dict] = []
    source = os.path.basename(docx_path)

    if output_path is None:
        base, ext = os.path.splitext(docx_path)
        output_path = base + '_fixed' + ext

    # Parse using low-level XML for analysis
    root = parse_docx_xml(docx_path)
    if root is None:
        return {
            "pass": False, "source": source, "check_type": "fix-docx",
            "action": "failed", "reason": "Failed to parse DOCX file",
            "heading_count": 0, "toc_entries_before": 0, "toc_entries_after": 0,
            "output": output_path,
            "errors": [make_item("PARSE_ERROR", "Failed to parse DOCX", "error")],
            "warnings": [], "info": []
        }

    body = root.find(_w('body'))
    if body is None:
        return {
            "pass": False, "source": source, "check_type": "fix-docx",
            "action": "failed", "reason": "Document body is empty",
            "heading_count": 0, "toc_entries_before": 0, "toc_entries_after": 0,
            "output": output_path,
            "errors": [make_item("EMPTY_BODY", "Document body is empty", "error")],
            "warnings": [], "info": []
        }

    paragraphs = [e for e in body if e.tag == _w('p')]

    # Extract headings
    headings = []  # list of (para_index_in_body, text, level)
    body_children = list(body)
    para_to_body_idx = {}  # map paragraph element to body child index
    pi = 0
    caption_filter = re.compile(r'^[表图]\s*\d')
    for ci, child in enumerate(body_children):
        if child.tag == _w('p'):
            para_to_body_idx[id(child)] = ci
            style = get_paragraph_style(child)
            if is_any_heading_style(style):
                text = get_paragraph_text(child).strip()
                level = _get_heading_level(style)
                if text and level > 0:
                    # Skip table/figure captions styled as headings
                    if caption_filter.match(text):
                        continue
                    headings.append((ci, text, level))
            pi += 1

    heading_count = len(headings)
    heading_texts = [h[1] for h in headings]

    # Get TOC info
    toc_info = find_toc_field_boundaries_v2(root)
    has_toc = toc_info['has_toc']
    toc_entry_texts = toc_info['toc_entry_texts']
    toc_entries_before = len(toc_entry_texts)

    # Also check for SDT-based TOC
    sdt_indices = _find_toc_sdt_indices(body)
    has_sdt_toc = len(sdt_indices) > 0

    has_any_toc = has_toc or has_sdt_toc

    # If SDT TOC, extract text from it for analysis
    if has_sdt_toc and not toc_entry_texts:
        for si in sdt_indices:
            sdt_elem = body_children[si]
            for t in sdt_elem.findall('.//' + _w('t')):
                if t.text and t.text.strip():
                    toc_entry_texts.append(t.text.strip())
        toc_entries_before = len(toc_entry_texts)

    # ---- Decision logic ----

    # Case 1: No TOC exists
    if not has_any_toc:
        if heading_count < 3:
            return {
                "pass": True, "source": source, "check_type": "fix-docx",
                "action": "no_toc_needed",
                "reason": f"Document has only {heading_count} headings, no TOC needed",
                "heading_count": heading_count,
                "toc_entries_before": 0, "toc_entries_after": 0,
                "output": output_path,
                "errors": [], "warnings": [],
                "info": [f"Document has {heading_count} headings (< 3), no TOC needed"]
            }
        else:
            # Need to generate TOC
            info_list.append(f"No TOC found, generating new TOC with {heading_count} entries")
            need_fix = True
            fix_reason = "no_toc"
            toc_insert_body_idx = None  # Will determine below
    else:
        # Case 2 & 3: TOC exists, check if it's stale/placeholder
        need_fix = False
        fix_reason = ""

        # Check for empty TOC
        non_empty_entries = [t for t in toc_entry_texts if t.strip()]
        if not non_empty_entries:
            need_fix = True
            fix_reason = "empty_toc"
            info_list.append("TOC exists but has no text content (uninitialized)")
        else:
            # Language mismatch check
            heading_lang = _detect_language(heading_texts)
            toc_lang = _detect_language(non_empty_entries)

            if heading_lang != toc_lang and heading_count >= 3:
                need_fix = True
                fix_reason = "language_mismatch"
                info_list.append(
                    f"Deleted stale TOC with {toc_entries_before} "
                    f"{'English' if toc_lang == 'en' else 'Chinese'} placeholder entries"
                )

            # Count mismatch check (>50% difference)
            if not need_fix and heading_count > 0:
                diff = abs(heading_count - toc_entries_before)
                if diff / heading_count > 0.5:
                    need_fix = True
                    fix_reason = "count_mismatch"
                    info_list.append(
                        f"TOC has {toc_entries_before} entries but document has "
                        f"{heading_count} headings (>{50}% drift)"
                    )

        if not need_fix:
            # TOC looks OK
            return {
                "pass": True, "source": source, "check_type": "fix-docx",
                "action": "skipped",
                "reason": "TOC appears to be up-to-date",
                "heading_count": heading_count,
                "toc_entries_before": toc_entries_before,
                "toc_entries_after": toc_entries_before,
                "output": output_path,
                "errors": [], "warnings": [],
                "info": ["TOC entries and headings are consistent, no fix needed"]
            }

    # ---- Perform the fix using python-docx ----
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        return {
            "pass": False, "source": source, "check_type": "fix-docx",
            "action": "failed", "reason": f"Failed to open DOCX with python-docx: {str(e)[:200]}",
            "heading_count": heading_count,
            "toc_entries_before": toc_entries_before, "toc_entries_after": 0,
            "output": output_path,
            "errors": [make_item("OPEN_ERROR", f"Failed to open: {str(e)[:200]}", "error")],
            "warnings": [], "info": []
        }

    doc_body = doc.element.body
    doc_children = list(doc_body)

    # Determine language for TOC title
    content_lang = _detect_language(heading_texts)
    toc_title = "目  录" if content_lang == 'zh' else "Table of Contents"

    # Re-extract headings from the python-docx document for consistency
    doc_headings = []  # (element_index, text, level)
    # Pattern to filter out table/figure captions styled as headings
    caption_re = re.compile(r'^[表图]\s*\d')
    for ci, child in enumerate(doc_children):
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'))
                    if is_any_heading_style(style_val):
                        text_parts = []
                        for t in child.findall('.//' + qn('w:t')):
                            if t.text:
                                text_parts.append(t.text)
                        text = ''.join(text_parts).strip()
                        level = _get_heading_level(style_val)
                        if text and level > 0:
                            # Skip table/figure captions (e.g. "表 1：xxx", "图 2：xxx")
                            if caption_re.match(text):
                                continue
                            doc_headings.append((ci, text, level))

    if not doc_headings:
        return {
            "pass": True, "source": source, "check_type": "fix-docx",
            "action": "no_toc_needed",
            "reason": "No headings found in document after re-parse",
            "heading_count": 0,
            "toc_entries_before": toc_entries_before, "toc_entries_after": 0,
            "output": output_path,
            "errors": [], "warnings": [],
            "info": ["No headings found, skipping TOC generation"]
        }

    # Step 1: Remove existing TOC (SDT or field code range)
    insert_before_idx = None

    # Remove SDT-based TOC
    sdt_removed = False
    for child in list(doc_body):
        if child.tag == qn('w:sdt'):
            is_toc_sdt = False
            for instr in child.findall('.//' + qn('w:instrText')):
                if instr.text and 'TOC' in instr.text.upper():
                    is_toc_sdt = True
                    break
            if not is_toc_sdt:
                sdtPr = child.find(qn('w:sdtPr'))
                if sdtPr is not None:
                    alias = sdtPr.find(qn('w:alias'))
                    if alias is not None and alias.get(qn('w:val'), '').upper() in ('TOC', '目录'):
                        is_toc_sdt = True
                    docPartObj = sdtPr.find(qn('w:docPartObj'))
                    if docPartObj is not None:
                        dpg = docPartObj.find(qn('w:docPartGallery'))
                        if dpg is not None and 'toc' in dpg.get(qn('w:val'), '').lower():
                            is_toc_sdt = True
            if is_toc_sdt:
                # Record position
                insert_before_idx = list(doc_body).index(child)
                doc_body.remove(child)
                sdt_removed = True

    # Remove field code TOC (non-SDT)
    if not sdt_removed and has_toc:
        # Find and remove paragraphs that are part of the TOC field
        doc_children_fresh = list(doc_body)
        # Use similar logic to find_toc_field_boundaries_v2 but on python-docx elements
        in_toc = False
        toc_depth = None
        depth = 0
        toc_paras_to_remove = []
        field_begin_idx = None

        for ci, child in enumerate(doc_children_fresh):
            if child.tag != qn('w:p'):
                continue
            for run in child.findall('.//' + qn('w:r')):
                instr = run.find(qn('w:instrText'))
                if instr is not None and instr.text and 'TOC' in instr.text.upper():
                    in_toc = True

                fldChar = run.find(qn('w:fldChar'))
                if fldChar is not None:
                    fld_type = fldChar.get(qn('w:fldCharType'), '')
                    if fld_type == 'begin':
                        depth += 1
                        if in_toc and field_begin_idx is None:
                            field_begin_idx = ci
                            toc_depth = depth
                    elif fld_type == 'end':
                        if in_toc and depth == toc_depth:
                            # Mark all paragraphs from begin to end for removal
                            if field_begin_idx is not None:
                                for ri in range(field_begin_idx, ci + 1):
                                    toc_paras_to_remove.append(doc_children_fresh[ri])
                            in_toc = False
                        depth = max(0, depth - 1)

        if toc_paras_to_remove:
            insert_before_idx = list(doc_body).index(toc_paras_to_remove[0])
            for p in toc_paras_to_remove:
                try:
                    doc_body.remove(p)
                except ValueError:
                    pass

    # Step 2: Determine insertion point
    if insert_before_idx is None:
        # No existing TOC was removed — find the right place to insert
        doc_children_now = list(doc_body)
        first_heading_idx = None
        for ci, child in enumerate(doc_children_now):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        sv = pStyle.get(qn('w:val'))
                        if is_any_heading_style(sv):
                            first_heading_idx = ci
                            break
        if first_heading_idx is not None:
            insert_before_idx = first_heading_idx
        else:
            insert_before_idx = 0

    # Step 3: Build TOC paragraphs as OxmlElements and insert them

    def _make_toc_paragraph(text: str, level: int, lang: str, page_num: str = '1', bookmark_name: str = '') -> Any:
        """Create a TOC entry paragraph with HYPERLINK + PAGEREF for clickable links and auto page numbers."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')

        # TOC style
        toc_style = OxmlElement('w:pStyle')
        toc_style.set(qn('w:val'), f'TOC{level}' if level <= 3 else 'TOC3')
        pPr.append(toc_style)

        # Indentation based on level
        if level >= 2:
            ind = OxmlElement('w:ind')
            indent_twips = (level - 1) * 420
            ind.set(qn('w:left'), str(indent_twips))
            pPr.append(ind)

        # Right-aligned tab stop with dot leader at 9026 twips (~15.9cm)
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9026')
        tabs.append(tab)
        pPr.append(tabs)

        # Line spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '120')
        spacing.set(qn('w:after'), '60')
        pPr.append(spacing)

        p.append(pPr)

        if bookmark_name:
            # Wrap everything in a hyperlink element pointing to the bookmark
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), bookmark_name)
            hyperlink.set(qn('w:history'), '1')

            # --- Run 1: heading text ---
            r = OxmlElement('w:r')
            rPr_r = OxmlElement('w:rPr')
            # Style as hyperlink (blue, underline optional)
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr_r.append(rStyle)
            sz2 = OxmlElement('w:sz')
            szCs2 = OxmlElement('w:szCs')
            if level == 1:
                sz2.set(qn('w:val'), '28')
                szCs2.set(qn('w:val'), '28')
                b2 = OxmlElement('w:b')
                rPr_r.append(b2)
            elif level == 2:
                sz2.set(qn('w:val'), '24')
                szCs2.set(qn('w:val'), '24')
            else:
                sz2.set(qn('w:val'), '22')
                szCs2.set(qn('w:val'), '22')
            rPr_r.append(sz2)
            rPr_r.append(szCs2)
            r.append(rPr_r)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text
            r.append(t)
            hyperlink.append(r)

            # --- Run 2: tab ---
            r_tab = OxmlElement('w:r')
            tab_elem = OxmlElement('w:tab')
            r_tab.append(tab_elem)
            hyperlink.append(r_tab)

            # --- Run 3: PAGEREF field code for auto page number ---
            # fldChar begin
            r_begin = OxmlElement('w:r')
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            r_begin.append(fldChar_begin)
            hyperlink.append(r_begin)

            # instrText: PAGEREF bookmark_name \h
            r_instr = OxmlElement('w:r')
            instrText = OxmlElement('w:instrText')
            instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            instrText.text = f' PAGEREF {bookmark_name} \\h '
            r_instr.append(instrText)
            hyperlink.append(r_instr)

            # fldChar separate
            r_sep = OxmlElement('w:r')
            fldChar_sep = OxmlElement('w:fldChar')
            fldChar_sep.set(qn('w:fldCharType'), 'separate')
            r_sep.append(fldChar_sep)
            hyperlink.append(r_sep)

            # Page number placeholder text
            r_page = OxmlElement('w:r')
            rPr_page = OxmlElement('w:rPr')
            noProof = OxmlElement('w:noProof')
            rPr_page.append(noProof)
            r_page.append(rPr_page)
            t_page = OxmlElement('w:t')
            t_page.text = str(page_num)
            r_page.append(t_page)
            hyperlink.append(r_page)

            # fldChar end
            r_end = OxmlElement('w:r')
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            r_end.append(fldChar_end)
            hyperlink.append(r_end)

            p.append(hyperlink)
        else:
            # Fallback: plain text without hyperlink (same as before)
            r = OxmlElement('w:r')
            rPr_r = OxmlElement('w:rPr')
            sz2 = OxmlElement('w:sz')
            szCs2 = OxmlElement('w:szCs')
            if level == 1:
                sz2.set(qn('w:val'), '28')
                szCs2.set(qn('w:val'), '28')
                b2 = OxmlElement('w:b')
                rPr_r.append(b2)
            elif level == 2:
                sz2.set(qn('w:val'), '24')
                szCs2.set(qn('w:val'), '24')
            else:
                sz2.set(qn('w:val'), '22')
                szCs2.set(qn('w:val'), '22')
            rPr_r.append(sz2)
            rPr_r.append(szCs2)
            r.append(rPr_r)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text
            r.append(t)
            p.append(r)

            r_tab = OxmlElement('w:r')
            tab_elem = OxmlElement('w:tab')
            r_tab.append(tab_elem)
            p.append(r_tab)

            r_page = OxmlElement('w:r')
            t_page = OxmlElement('w:t')
            t_page.text = str(page_num)
            r_page.append(t_page)
            p.append(r_page)

        return p

    def _make_toc_title(title_text: str) -> Any:
        """Create the TOC title paragraph (centered, 18pt, bold)."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')

        # Center alignment
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

        # Spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '200')
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)

        # Run properties
        rPr_p = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr_p.append(b)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '36')  # 18pt = 36 half-points
        rPr_p.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '36')
        rPr_p.append(szCs)
        pPr.append(rPr_p)

        p.append(pPr)

        # Run with text
        r = OxmlElement('w:r')
        rPr_r = OxmlElement('w:rPr')
        b2 = OxmlElement('w:b')
        rPr_r.append(b2)
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '36')
        rPr_r.append(sz2)
        szCs2 = OxmlElement('w:szCs')
        szCs2.set(qn('w:val'), '36')
        rPr_r.append(szCs2)
        r.append(rPr_r)

        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = title_text
        r.append(t)
        p.append(r)

        return p

    def _make_page_break() -> Any:
        """Create a paragraph with a page break."""
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        return p

    # Build the TOC as an SDT (Structured Document Tag) wrapping a TOC field
    # This ensures check-docx recognizes it and fix-docx can detect/replace it

    def _build_toc_sdt(title_text: str, heading_entries: list, lang: str) -> Any:
        """Build a complete SDT element containing a TOC field with entries."""
        sdt = OxmlElement('w:sdt')

        # SDT properties
        sdtPr = OxmlElement('w:sdtPr')
        alias = OxmlElement('w:alias')
        alias.set(qn('w:val'), 'TOC')
        sdtPr.append(alias)

        # docPartObj with TOC gallery
        docPartObj = OxmlElement('w:docPartObj')
        docPartGallery = OxmlElement('w:docPartGallery')
        docPartGallery.set(qn('w:val'), 'Table of Contents')
        docPartObj.append(docPartGallery)
        docPartUnique = OxmlElement('w:docPartUnique')
        docPartObj.append(docPartUnique)
        sdtPr.append(docPartObj)

        sdt.append(sdtPr)

        # SDT content
        sdtContent = OxmlElement('w:sdtContent')

        # Title paragraph
        sdtContent.append(_make_toc_title(title_text))

        # Field begin paragraph
        p_begin = OxmlElement('w:p')
        r_begin = OxmlElement('w:r')
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        r_begin.append(fldChar_begin)
        p_begin.append(r_begin)
        r_instr = OxmlElement('w:r')
        instrText = OxmlElement('w:instrText')
        instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        r_instr.append(instrText)
        p_begin.append(r_instr)
        r_sep = OxmlElement('w:r')
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        r_sep.append(fldChar_sep)
        p_begin.append(r_sep)
        sdtContent.append(p_begin)

        # TOC entry paragraphs — estimate page numbers based on heading position
        toc_entries = [(i, h_text, h_level) for i, (_, h_text, h_level) in enumerate(heading_entries) if h_level <= 3]
        total_headings = len(toc_entries)
        # TOC itself takes ~2 pages; cover takes ~1 page
        toc_offset = 3  # cover + TOC pages

        # Count body children in the original document to estimate total pages
        # Rough heuristic: ~40 paragraphs per page for typical documents
        doc_body_count = len(list(doc_body))
        estimated_total_pages = max(toc_offset + 1, doc_body_count // 40 + toc_offset)

        # Map each heading to its position ratio in the document
        # Also generate bookmark names for HYPERLINK + PAGEREF
        bookmark_names = []
        for seq, (idx, h_text, h_level) in enumerate(toc_entries):
            # Generate a unique bookmark name for each heading
            bm_name = f'_Toc{100000 + seq}'
            bookmark_names.append(bm_name)

            # Use the heading's body child index to estimate position
            h_body_idx = heading_entries[idx][0] if idx < len(heading_entries) else 0
            if doc_body_count > 0:
                position_ratio = h_body_idx / doc_body_count
                est_page = toc_offset + max(0, int(position_ratio * (estimated_total_pages - toc_offset)))
            else:
                est_page = toc_offset + seq
            est_page = max(toc_offset, est_page)  # never less than toc_offset
            sdtContent.append(_make_toc_paragraph(h_text, h_level, lang, str(est_page), bm_name))

        # Field end paragraph
        p_end = OxmlElement('w:p')
        r_end = OxmlElement('w:r')
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        r_end.append(fldChar_end)
        p_end.append(r_end)
        sdtContent.append(p_end)

        sdt.append(sdtContent)

        # Build bookmark mapping: list of (heading_body_index, bookmark_name)
        bm_mapping = []
        for seq, (idx, h_text, h_level) in enumerate(toc_entries):
            h_body_idx = heading_entries[idx][0] if idx < len(heading_entries) else 0
            bm_mapping.append((h_body_idx, bookmark_names[seq]))

        return sdt, bm_mapping

    # Build TOC SDT and page break
    toc_sdt, bookmark_mapping = _build_toc_sdt(toc_title, doc_headings, content_lang)
    page_break = _make_page_break()

    # Insert TOC elements at the determined position
    ref_children = list(doc_body)
    # Clamp insert_before_idx
    if insert_before_idx >= len(ref_children):
        doc_body.append(toc_sdt)
        doc_body.append(page_break)
    else:
        ref_element = ref_children[insert_before_idx]
        ref_element.addprevious(toc_sdt)
        ref_element.addprevious(page_break)

    # Add bookmarks to heading paragraphs so PAGEREF and HYPERLINK can resolve
    body_children = list(doc_body)
    bm_id_start = 10  # bookmark IDs must be unique integers in the document
    for body_idx, bm_name in bookmark_mapping:
        if body_idx < len(body_children):
            heading_para = body_children[body_idx]
            # Insert bookmarkStart before first run, bookmarkEnd after last run
            bm_start = OxmlElement('w:bookmarkStart')
            bm_start.set(qn('w:id'), str(bm_id_start))
            bm_start.set(qn('w:name'), bm_name)

            bm_end = OxmlElement('w:bookmarkEnd')
            bm_end.set(qn('w:id'), str(bm_id_start))

            # Insert at beginning and end of the paragraph
            heading_para.insert(0, bm_start)
            heading_para.append(bm_end)

            bm_id_start += 1

    # Save
    try:
        doc.save(output_path)
    except Exception as e:
        return {
            "pass": False, "source": source, "check_type": "fix-docx",
            "action": "failed", "reason": f"Failed to save: {str(e)[:200]}",
            "heading_count": heading_count,
            "toc_entries_before": toc_entries_before, "toc_entries_after": 0,
            "output": output_path,
            "errors": [make_item("SAVE_ERROR", f"Failed to save: {str(e)[:200]}", "error")],
            "warnings": [], "info": []
        }

    toc_entries_after = sum(1 for _, _, l in doc_headings if l <= 3)
    info_list.append(f"Generated new TOC with {toc_entries_after} entries")

    return {
        "pass": True, "source": source, "check_type": "fix-docx",
        "action": "fixed",
        "reason": fix_reason,
        "heading_count": heading_count,
        "toc_entries_before": toc_entries_before,
        "toc_entries_after": toc_entries_after,
        "output": output_path,
        "errors": [], "warnings": [],
        "info": info_list
    }


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
def fix_docx_accurate_pages(fixed_docx_path: str, pass1_pdf_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
    """Update TOC page numbers in a fix-docx output using actual page positions from a PDF.

    Two-pass approach:
      Pass 1: Convert the DOCX (without TOC fix or with estimated pages) to PDF
      Pass 2: Read actual heading positions from PDF, update PAGEREF placeholder text

    Args:
        fixed_docx_path: Path to the DOCX after fix-docx (has PAGEREF fields with estimated pages)
        pass1_pdf_path: Path to a PDF converted from the ORIGINAL docx (without TOC)
        output_path: Where to save the updated DOCX (defaults to overwrite fixed_docx_path)
    """
    import zipfile as zf_mod
    import tempfile
    import shutil

    try:
        import pdfplumber
    except ImportError:
        return {"pass": False, "error": "pdfplumber not installed — cannot extract page positions"}

    try:
        from docx import Document
        from docx.oxml.ns import qn as docx_qn
    except ImportError:
        return {"pass": False, "error": "python-docx not installed"}

    try:
        from lxml import etree
    except ImportError:
        return {"pass": False, "error": "lxml not installed"}

    if output_path is None:
        output_path = fixed_docx_path

    source = os.path.basename(fixed_docx_path)

    # --- Step 1: Extract headings from the fixed DOCX ---
    doc = Document(fixed_docx_path)
    headings = []
    caption_pattern = re.compile(r'^[表图]\s*\d')
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if style_name.startswith('Heading'):
            m = re.match(r'Heading\s*(\d+)', style_name)
            if m:
                text = p.text.strip()
                if text and not caption_pattern.match(text):
                    headings.append({'text': text, 'level': int(m.group(1))})

    if not headings:
        return {"pass": True, "source": source, "info": "No headings found, nothing to update"}

    # --- Step 2: Find actual page positions in pass1 PDF ---
    pdf = pdfplumber.open(pass1_pdf_path)
    total_pdf_pages = len(pdf.pages)

    page_texts = []
    for i in range(total_pdf_pages):
        page_texts.append(pdf.pages[i].extract_text() or '')

    heading_pages_pass1: Dict[str, int] = {}
    for h in headings:
        for page_num, pt in enumerate(page_texts):
            if h['text'] in pt:
                heading_pages_pass1[h['text']] = page_num + 1  # 1-indexed
                break
    pdf.close()

    # --- Step 3: Calculate offset ---
    # Instead of estimating TOC page count, calculate actual offset by comparing
    # where the first heading appears in pass1 vs where it should appear after TOC insertion.
    # The offset = (number of pages TOC adds) which depends on entry count and formatting.
    toc_entry_count = sum(1 for h in headings if h['level'] <= 3)

    # Better estimate: ~15 entries per page for CJK text with leader dots
    toc_pages = max(1, (toc_entry_count + 14) // 15)

    # Additional offset for the page break after TOC
    # Check if the original DOCX already had a TOC (pass1 already includes TOC space)
    # by looking at whether the first heading is on page 1-2 (no TOC) or later (has TOC)
    first_heading_page = min(heading_pages_pass1.values()) if heading_pages_pass1 else 1

    if first_heading_page <= 2:
        # Pass1 has no significant TOC content, so we need full offset
        offset = toc_pages + 1  # +1 for page break after TOC
    else:
        # Pass1 already has some TOC pages, smaller offset needed
        offset = max(0, toc_pages - (first_heading_page - 2))

    heading_page_map: Dict[str, int] = {}
    for h_text, orig_page in heading_pages_pass1.items():
        heading_page_map[h_text] = orig_page + offset

    # --- Step 4: Update PAGEREF placeholder text in the DOCX XML ---
    with zf_mod.ZipFile(fixed_docx_path, 'r') as zf:
        doc_xml = zf.read('word/document.xml')

    root = etree.fromstring(doc_xml)
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    w_ns = nsmap['w']

    all_runs = root.findall(f'.//{{{w_ns}}}r')
    updates = 0
    i = 0
    while i < len(all_runs):
        r = all_runs[i]
        fld = r.find(f'{{{w_ns}}}fldChar')
        if fld is not None and fld.get(f'{{{w_ns}}}fldCharType') == 'begin':
            if i + 1 < len(all_runs):
                instr_r = all_runs[i + 1]
                instr_t = instr_r.find(f'{{{w_ns}}}instrText')
                if instr_t is not None and instr_t.text and 'PAGEREF' in instr_t.text:
                    # Find the hyperlink parent to get heading text
                    hyperlink = r.getparent()
                    if hyperlink is not None:
                        text_runs = hyperlink.findall(f'.//{{{w_ns}}}t')
                        # Find the 'separate' then the page number text
                        for j in range(i + 2, min(i + 5, len(all_runs))):
                            sep_fld = all_runs[j].find(f'{{{w_ns}}}fldChar')
                            if sep_fld is not None and sep_fld.get(f'{{{w_ns}}}fldCharType') == 'separate':
                                if j + 1 < len(all_runs):
                                    page_t = all_runs[j + 1].find(f'{{{w_ns}}}t')
                                    if page_t is not None:
                                        # Get heading text from the hyperlink
                                        heading_text = ''
                                        for tr in text_runs:
                                            if tr.text and tr != page_t:
                                                heading_text += tr.text
                                        heading_text = heading_text.strip()

                                        correct_page = heading_page_map.get(heading_text)
                                        if correct_page:
                                            page_t.text = str(correct_page)
                                            updates += 1
                                break
        i += 1

    # --- Step 5: Save updated DOCX ---
    with tempfile.TemporaryDirectory() as tmpdir:
        with zf_mod.ZipFile(fixed_docx_path, 'r') as zf:
            zf.extractall(tmpdir)

        doc_xml_path = os.path.join(tmpdir, 'word', 'document.xml')
        with open(doc_xml_path, 'wb') as f:
            f.write(etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True))

        with zf_mod.ZipFile(output_path, 'w', zf_mod.ZIP_DEFLATED) as zf:
            for dirpath, dirnames, filenames in os.walk(tmpdir):
                for fn in filenames:
                    full_path = os.path.join(dirpath, fn)
                    arcname = os.path.relpath(full_path, tmpdir)
                    zf.write(full_path, arcname)

    return {
        "pass": True,
        "source": source,
        "check_type": "fix-pages",
        "pages_updated": updates,
        "total_headings": len(headings),
        "toc_pages_estimated": toc_pages,
        "offset_applied": offset,
        "output": output_path,
    }


def print_usage():
    """Print usage information to stderr."""
    print("Usage:", file=sys.stderr)
    print("  toc_validate.py check-docx <file.docx>", file=sys.stderr)
    print("  toc_validate.py check-pdf  <file.pdf>", file=sys.stderr)
    print("  toc_validate.py check-conversion <input.docx> <output.pdf>",
          file=sys.stderr)
    print("  toc_validate.py fix-docx <input.docx> [-o output.docx]",
          file=sys.stderr)
    print("  toc_validate.py fix-pages <fixed.docx> <pass1.pdf> [-o output.docx]",
          file=sys.stderr)
    print("", file=sys.stderr)
    print("fix-pages: 2-pass page number correction. Requires a PDF converted", file=sys.stderr)
    print("           from the ORIGINAL docx (without TOC) as reference.", file=sys.stderr)


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print_usage()
        sys.exit(2)

    command = sys.argv[1].lower()

    try:
        if command == 'check-docx':
            if len(sys.argv) < 3:
                print("Error: Missing DOCX file path.", file=sys.stderr)
                print_usage()
                sys.exit(2)
            docx_path = sys.argv[2]
            if not os.path.isfile(docx_path):
                print(f"Error: File not found: {docx_path}", file=sys.stderr)
                sys.exit(2)
            result = check_docx(docx_path)

        elif command == 'check-pdf':
            if len(sys.argv) < 3:
                print("Error: Missing PDF file path.", file=sys.stderr)
                print_usage()
                sys.exit(2)
            pdf_path = sys.argv[2]
            if not os.path.isfile(pdf_path):
                print(f"Error: File not found: {pdf_path}", file=sys.stderr)
                sys.exit(2)
            result = check_pdf(pdf_path)

        elif command == 'check-conversion':
            if len(sys.argv) < 4:
                print("Error: Missing file paths. Need both DOCX and PDF.",
                      file=sys.stderr)
                print_usage()
                sys.exit(2)
            docx_path = sys.argv[2]
            pdf_path = sys.argv[3]
            if not os.path.isfile(docx_path):
                print(f"Error: File not found: {docx_path}", file=sys.stderr)
                sys.exit(2)
            if not os.path.isfile(pdf_path):
                print(f"Error: File not found: {pdf_path}", file=sys.stderr)
                sys.exit(2)
            result = check_conversion(docx_path, pdf_path)

        elif command == 'fix-docx':
            if len(sys.argv) < 3:
                print("Error: Missing DOCX file path.", file=sys.stderr)
                print_usage()
                sys.exit(2)
            docx_path = sys.argv[2]
            if not os.path.isfile(docx_path):
                print(f"Error: File not found: {docx_path}", file=sys.stderr)
                sys.exit(2)
            # Parse optional -o flag
            output_path = None
            if '-o' in sys.argv:
                o_idx = sys.argv.index('-o')
                if o_idx + 1 < len(sys.argv):
                    output_path = sys.argv[o_idx + 1]
                else:
                    print("Error: -o flag requires an output path.",
                          file=sys.stderr)
                    sys.exit(2)
            result = fix_docx(docx_path, output_path)

        elif command == 'fix-pages':
            if len(sys.argv) < 4:
                print("Error: Need both fixed DOCX and pass1 PDF paths.", file=sys.stderr)
                print_usage()
                sys.exit(2)
            fixed_docx = sys.argv[2]
            pass1_pdf = sys.argv[3]
            if not os.path.isfile(fixed_docx):
                print(f"Error: File not found: {fixed_docx}", file=sys.stderr)
                sys.exit(2)
            if not os.path.isfile(pass1_pdf):
                print(f"Error: File not found: {pass1_pdf}", file=sys.stderr)
                sys.exit(2)
            output_path = None
            if '-o' in sys.argv:
                o_idx = sys.argv.index('-o')
                if o_idx + 1 < len(sys.argv):
                    output_path = sys.argv[o_idx + 1]
            result = fix_docx_accurate_pages(fixed_docx, pass1_pdf, output_path)

        else:
            print(f"Error: Unknown command '{command}'", file=sys.stderr)
            print_usage()
            sys.exit(2)

        # Output JSON to stdout
        print(json.dumps(result, ensure_ascii=False, indent=2))

        # Exit code: 0=pass, 1=fail
        sys.exit(0 if result['pass'] else 1)

    except Exception as e:
        # Unexpected error — output JSON error and exit 2
        error_result = {
            "pass": False,
            "source": sys.argv[2] if len(sys.argv) > 2 else "unknown",
            "check_type": command.replace('check-', '') + '-toc'
                          if command.startswith('check-') else 'unknown',
            "errors": [make_item("SCRIPT_ERROR",
                                 f"Unexpected error: {str(e)[:200]}",
                                 "error")],
            "warnings": [],
            "info": [],
        }
        print(json.dumps(error_result, ensure_ascii=False, indent=2))
        sys.exit(2)


if __name__ == '__main__':
    main()
